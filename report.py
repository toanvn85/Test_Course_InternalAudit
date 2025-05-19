import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import io
import base64
from datetime import datetime
import numpy as np
import json
import sys
import traceback
import os

from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Thêm vào đầu file - Thay đổi sang fpdf2 thay vì fpdf
from fpdf import FPDF
# Thêm thư viện để hỗ trợ Unicode
import pkg_resources

from database_helper import get_supabase_client

# Thêm vào đầu chương trình 
import subprocess
import sys


# Thêm vào đầu file
import sys
import os
import subprocess

# Kiểm tra và cài đặt pip nếu cần
def install_pip():
    try:
        print("Đang cài đặt pip...")
        # Tải get-pip.py
        import urllib.request
        urllib.request.urlretrieve("https://bootstrap.pypa.io/get-pip.py", "get-pip.py")
        
        # Chạy get-pip.py
        subprocess.check_call([sys.executable, "get-pip.py"])
        print("Đã cài đặt pip thành công!")
        return True
    except Exception as e:
        print(f"Lỗi khi cài đặt pip: {str(e)}")
        return False

# Kiểm tra và cài đặt FPDF2 nếu cần
try:
    import pkg_resources
    fpdf_version = pkg_resources.get_distribution("fpdf").version
    if not fpdf_version.startswith("2."):
        print("Đang cài đặt FPDF2 cho hỗ trợ Unicode...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "--upgrade", "fpdf2"])
        print("Đã cài đặt FPDF2 thành công!")
except Exception as e:
    print(f"Lỗi khi cài đặt FPDF2: {str(e)}")

# Thêm vào đầu file sau khi import các thư viện
# Kiểm tra phiên bản FPDF
try:
    fpdf_pkg = pkg_resources.get_distribution("fpdf")
    FPDF_VERSION = fpdf_pkg.version
    IS_FPDF2 = FPDF_VERSION.startswith("2.")
    if IS_FPDF2:
        print(f"Đang sử dụng FPDF2 (phiên bản {FPDF_VERSION})")
    else:
        print(f"Đang sử dụng FPDF1 (phiên bản {FPDF_VERSION})")
except Exception as e:
    print(f"Không thể xác định phiên bản FPDF: {str(e)}")
    IS_FPDF2 = False
    FPDF_VERSION = "unknown"

# Giả lập database_helper nếu không có
try:
    from database_helper import check_answer_correctness, get_all_questions, get_all_users, get_user_submissions
except ImportError:
    # Mock functions để tránh lỗi khi không có module
    def check_answer_correctness(user_ans, q):
        q_correct = q.get("correct", [])
        if isinstance(q_correct, str):
            try:
                q_correct = json.loads(q_correct)
            except:
                try:
                    q_correct = [int(x.strip()) for x in q_correct.split(",")]
                except:
                    q_correct = []
        
        # Chuyển user_ans và q_correct thành tập hợp để so sánh
        if isinstance(user_ans, list) and isinstance(q_correct, list):
            # Chuyển đổi các đáp án thành chuỗi để có thể so sánh
            user_ans_set = set(str(x) for x in user_ans)
            correct_set = set(str(x) for x in q_correct)
            return user_ans_set == correct_set
        return False
    
    def get_all_questions():
        return []
    
    def get_all_users(role=None):
        return []
    
    def get_user_submissions(email):
        return []

# Nhập các thư viện cho xuất file
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    # Hiển thị thông báo chỉ khi đang chạy trong Streamlit
    if 'streamlit' in sys.modules:
        st.warning("Module python-docx không được cài đặt. Tính năng xuất DOCX sẽ không hoạt động.")

# Sử dụng WD_ALIGN_PARAGRAPH nếu có thể, nếu không tạo class thay thế
try:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    class WD_ALIGN_PARAGRAPH:
        CENTER = 1
        RIGHT = 2
        LEFT = 0

# Hỗ trợ xuất PDF với reportlab nếu cần
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
except ImportError:
    # Hiển thị thông báo chỉ khi đang chạy trong Streamlit
    if 'streamlit' in sys.modules:
        st.warning("Module reportlab không được cài đặt. Tính năng xuất PDF sẽ bị hạn chế.")

# Hàm kiểm tra cài đặt và phiên bản của FPDF
def check_fpdf_installed():
    """Kiểm tra cài đặt và phiên bản của FPDF"""
    try:
        # Kiểm tra phiên bản của fpdf
        fpdf_pkg = pkg_resources.get_distribution("fpdf")
        is_fpdf2 = fpdf_pkg.version.startswith("2.")
        
        if is_fpdf2:
            st.success(f"FPDF2 đã được cài đặt, phiên bản: {fpdf_pkg.version}")
            st.info("FPDF2 có hỗ trợ Unicode tốt hơn. Xuất PDF với tiếng Việt sẽ hoạt động tốt.")
        else:
            st.warning(f"FPDF1 đã được cài đặt, phiên bản: {fpdf_pkg.version}. Để hỗ trợ tiếng Việt tốt hơn, hãy cài đặt FPDF2 bằng lệnh: pip install fpdf2")
            st.info("Báo cáo PDF sẽ không hiển thị đúng tiếng Việt có dấu. Khuyến nghị sử dụng định dạng DOCX.")
        return True
    except pkg_resources.DistributionNotFound:
        st.error("FPDF chưa được cài đặt. Hãy cài đặt bằng lệnh: pip install fpdf2")
        return False
    except Exception as e:
        st.error(f"Lỗi khi kiểm tra FPDF: {str(e)}")
        return False

# Chuẩn bị font tiếng Việt
def setup_vietnamese_fonts():
    """Cài đặt và đăng ký font cho tiếng Việt"""
    current_dir = os.path.dirname(os.path.abspath(__file__)) if '__file__' in globals() else os.getcwd()
    
    # Các đường dẫn có thể chứa font
    font_dirs = [
        os.path.join(current_dir, 'assets', 'fonts'),
        os.path.join(current_dir, 'fonts'),
        os.path.join(current_dir, 'assets'),
        current_dir,
        '/usr/share/fonts/truetype',
        '/usr/share/fonts/truetype/dejavu',
        '/usr/share/fonts/TTF',
        'C:\\Windows\\Fonts',
    ]
    
    # Các font cần tìm theo thứ tự ưu tiên
    font_files = [
        ('DejaVuSans', 'DejaVuSans.ttf'),
        ('DejaVuSans-Bold', 'DejaVuSans-Bold.ttf'),
        ('DejaVuSans-Oblique', 'DejaVuSans-Oblique.ttf'),
        ('Arial', 'arial.ttf'),
        ('Arial-Bold', 'arialbd.ttf'),
        ('Arial-Italic', 'ariali.ttf'),
        ('TimesNewRoman', 'times.ttf'),
        ('TimesNewRoman-Bold', 'timesbd.ttf'),
    ]
    
    registered_fonts = []
    
    # Tìm và đăng ký font
    for font_name, font_file in font_files:
        for font_dir in font_dirs:
            font_path = os.path.join(font_dir, font_file)
            if os.path.exists(font_path):
                try:
                    # Đăng ký font với reportlab nếu đã import
                    if 'pdfmetrics' in globals() and 'TTFont' in globals():
                        try:
                            pdfmetrics.registerFont(TTFont(font_name, font_path))
                            print(f"Đã đăng ký font {font_name} từ {font_path}")
                        except Exception as e:
                            print(f"Không thể đăng ký font {font_name} cho reportlab: {str(e)}")
                    
                    registered_fonts.append((font_name, font_path))
                    break
                except Exception as e:
                    print(f"Lỗi khi đăng ký font {font_name}: {str(e)}")
    
    # Thêm các font mặc định vào cuối danh sách nếu chưa có font nào được đăng ký
    if not registered_fonts:
        registered_fonts.append(('Helvetica', ''))
        registered_fonts.append(('Courier', ''))
    
    return registered_fonts

def format_date(date_value):
    """Định dạng ngày tháng từ nhiều kiểu dữ liệu khác nhau"""
    if not date_value:
        return "N/A"
    
    try:
        # Nếu là số nguyên (timestamp)
        if isinstance(date_value, (int, float)):
            return datetime.fromtimestamp(date_value).strftime("%d/%m/%Y")
        
        # Nếu là chuỗi ISO (từ Supabase)
        elif isinstance(date_value, str):
            try:
                # Thử parse chuỗi ISO
                dt = datetime.fromisoformat(date_value.replace('Z', '+00:00'))
                return dt.strftime("%d/%m/%Y")
            except:
                # Nếu không phải ISO, trả về nguyên bản
                return date_value
        
        # Nếu đã là đối tượng datetime
        elif isinstance(date_value, datetime):
            return date_value.strftime("%d/%m/%Y")
            
        # Các trường hợp khác, trả về dạng chuỗi
        else:
            return str(date_value)
    except Exception as e:
        print(f"Error formatting date: {e}, value type: {type(date_value)}, value: {date_value}")
        return "N/A"

def get_download_link_docx(buffer, filename, text):
    """Tạo link tải xuống cho file DOCX"""
    b64 = base64.b64encode(buffer.getvalue()).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">📥 {text}</a>'
    return href

def get_download_link_pdf(buffer, filename, text):
    """Tạo link tải xuống cho file PDF"""
    try:
        if buffer and hasattr(buffer, 'getvalue') and len(buffer.getvalue()) > 0:
            b64 = base64.b64encode(buffer.getvalue()).decode()
            href = f'<a href="data:application/pdf;base64,{b64}" download="{filename}">📥 {text}</a>'
            return href
        else:
            return f'<span style="color:red;">Không thể tạo link tải PDF. Vui lòng sử dụng định dạng DOCX thay thế.</span>'
    except Exception as e:
        print(f"Lỗi khi tạo link tải PDF: {str(e)}")
        return f'<span style="color:red;">Lỗi tạo link tải PDF: {str(e)}</span>'

def export_to_excel(dataframes, sheet_names, filename):
    """Tạo file Excel với nhiều sheet từ các DataFrame"""
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        for df, sheet_name in zip(dataframes, sheet_names):
            df.to_excel(writer, sheet_name=sheet_name, index=False)
    
    data = output.getvalue()
    b64 = base64.b64encode(data).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64}" download="{filename}">📥 {filename}</a>'
    return href

def dataframe_to_docx(df, title, filename):
    """Tạo file DOCX từ DataFrame"""
    try:
        doc = Document()
        
        # Thiết lập font chữ mặc định
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Thêm tiêu đề
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thêm thời gian xuất báo cáo
        time_paragraph = doc.add_paragraph(f"Thời gian xuất báo cáo: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        time_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Tạo bảng
        # Thêm một hàng cho tiêu đề cột
        table = doc.add_table(rows=1, cols=len(df.columns), style='Table Grid')
        
        # Thêm tiêu đề cột
        header_cells = table.rows[0].cells
        for i, col_name in enumerate(df.columns):
            header_cells[i].text = str(col_name)
            # Đặt kiểu cho tiêu đề
            for paragraph in header_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(str(col_name))
                run.bold = True
        
        # Thêm dữ liệu
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)
                # Căn giữa cho các ô
                for paragraph in row_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thêm chân trang
        doc.add_paragraph()
        footer = doc.add_paragraph("Hệ thống Khảo sát & Đánh giá")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Lưu tệp
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    except Exception as e:
        print(f"Lỗi khi tạo DOCX: {str(e)}")
        st.error(f"Không thể tạo file DOCX: {str(e)}")
        # Trả về buffer trống nếu lỗi
        buffer = io.BytesIO()
        buffer.seek(0)
        return buffer

class UNIOCDF_FPDF(FPDF):
    """Lớp PDF tùy chỉnh hỗ trợ Unicode đầy đủ"""
    def __init__(self, orientation='P', unit='mm', format='A4', title='Báo cáo'):
        super().__init__(orientation=orientation, unit=unit, format=format)
        self.title = title
        
        # Kiểm tra phiên bản FPDF
        try:
            fpdf_version = pkg_resources.get_distribution("fpdf").version
            self.is_fpdf2 = fpdf_version.startswith("2.")
        except:
            self.is_fpdf2 = False
            
        # Khắc phục lỗi tiếng Việt bằng cách thiết lập encode utf8 (chỉ cho FPDF2)
        if self.is_fpdf2:
            try:
                self.set_doc_option('core_fonts_encoding', 'utf-8')
            except:
                pass
        
        # Thêm các font có sẵn
        font_dirs = [
            os.path.dirname(os.path.abspath(__file__)) if '__file__' in globals() else os.getcwd(),
            'C:\\Windows\\Fonts',
        ]
        
        font_files = [
            ('Arial', 'arial.ttf'),
            ('Arial-Bold', 'arialbd.ttf'),
            ('Arial-Italic', 'ariali.ttf'),
        ]
        
        # Tìm và đăng ký font
        for font_name, font_file in font_files:
            for font_dir in font_dirs:
                font_path = os.path.join(font_dir, font_file)
                if os.path.exists(font_path):
                    try:
                        if self.is_fpdf2:
                            self.add_font(font_name, '', font_path, uni=True)
                        else:
                            self.add_font(font_name, '', font_path)
                        break
                    except Exception as e:
                        print(f"Lỗi khi đăng ký font {font_name}: {str(e)}")
        
    def header(self):
        # Font và tiêu đề
        try:
            self.set_font('Arial', 'B', 15)
        except:
            self.set_font('Helvetica', 'B', 15)
            
        # Tiêu đề ở giữa
        self.cell(0, 10, self.title, 0, 1, 'C')
        
        # Thời gian
        try:
            self.set_font('Arial', 'I', 8)
        except:
            self.set_font('Helvetica', 'I', 8)
            
        timestamp = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        self.cell(0, 5, f'Thoi gian xuat bao cao: {timestamp}', 0, 1, 'R')
        
        # Line break
        self.ln(5)
    
    def footer(self):
        # Vị trí cách đáy 15 mm
        self.set_y(-15)
        
        # Font
        try:
            self.set_font('Arial', 'I', 8)
        except:
            self.set_font('Helvetica', 'I', 8)
            
        # Số trang
        self.cell(0, 10, f'Trang {self.page_no()}/{self.alias_nb_pages()}', 0, 0, 'C')
        
        # Thêm chân trang hệ thống
        self.cell(0, 10, 'He thong kiểm tra Đánh giá viên nội bộ ISO 50001:2018 TUV', 0, 0, 'R')

# Tạo một instance FPDF có khả năng xử lý Unicode
def create_unicode_pdf(orientation='P', format='A4', title='Báo cáo'):
    """Tạo FPDF với hỗ trợ Unicode tốt hơn"""
    try:
        # Đăng ký font trước khi tạo PDF
        registered_fonts = setup_vietnamese_fonts()
        font_found = len(registered_fonts) > 0
        
        # Tạo PDF mới - Kiểm tra phiên bản FPDF
        try:
            fpdf_version = pkg_resources.get_distribution("fpdf").version
            is_fpdf2 = fpdf_version.startswith("2.")
        except:
            is_fpdf2 = False
            
        pdf = FPDF(orientation=orientation, unit='mm', format=format)
        
        # Thiết lập mã hóa UTF-8 cho FPDF2
        if is_fpdf2:
            try:
                pdf.set_doc_option('core_fonts_encoding', 'utf-8')
            except:
                pass
        
        # Thêm các font Unicode theo thứ tự ưu tiên
        if font_found:
            for font_name, font_path in registered_fonts:
                try:
                    # Đối với FPDF2, chúng ta sử dụng add_font với uni=True
                    # Đối với FPDF1, chúng ta không sử dụng tham số uni
                    if is_fpdf2:
                        pdf.add_font(font_name, '', font_path, uni=True)
                    else:
                        pdf.add_font(font_name, '', font_path)
                        
                    if font_name.endswith('-Bold'):
                        if is_fpdf2:
                            pdf.add_font(font_name.replace('-Bold', ''), 'B', font_path, uni=True)
                        else:
                            pdf.add_font(font_name.replace('-Bold', ''), 'B', font_path)
                    elif font_name.endswith('-Oblique') or font_name.endswith('-Italic'):
                        if is_fpdf2:
                            pdf.add_font(font_name.split('-')[0], 'I', font_path, uni=True)
                        else:
                            pdf.add_font(font_name.split('-')[0], 'I', font_path)
                except Exception as e:
                    print(f"Lỗi khi thêm font {font_name}: {str(e)}")
        else:
            # Nếu không tìm thấy font Unicode, sử dụng font mặc định
            print("Không tìm thấy font Unicode, sử dụng font mặc định")
        
        # Thiết lập các tùy chọn khác
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.alias_nb_pages()
        
        # Thiết lập tựa đề
        pdf.set_title(title)
        
        return pdf
    except Exception as e:
        print(f"Lỗi tạo PDF: {str(e)}")
        traceback.print_exc()
        
        # Phương án dự phòng - sử dụng FPDF cơ bản
        try:
            pdf = FPDF(orientation=orientation, format=format)
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.alias_nb_pages()
            return pdf
        except Exception as e2:
            print(f"Lỗi khi tạo PDF dự phòng: {str(e2)}")
            return None

def dataframe_to_pdf_fpdf(df, title, filename):
    """Tạo file PDF từ DataFrame sử dụng FPDF với xử lý lỗi Unicode"""
    buffer = io.BytesIO()
    try:
        # Xác định hướng trang dựa vào số lượng cột
        orientation = 'L' if len(df.columns) > 5 else 'P'
        
        # Khởi tạo đối tượng PDF
        pdf = FPDF(orientation=orientation, unit='mm', format='A4')
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        
        # Thêm tiêu đề - luôn dùng font mặc định
        pdf.set_font('Arial', 'B', 16)
        
        # Loại bỏ dấu tiếng Việt từ tiêu đề
        title_ascii = title.encode('ascii', 'ignore').decode('ascii')
        pdf.cell(0, 10, title_ascii, 0, 1, 'C')
        
        # Thêm thời gian báo cáo
        pdf.set_font('Arial', 'I', 10)
        timestamp = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        pdf.cell(0, 5, f'Thoi gian xuat bao cao: {timestamp}', 0, 1, 'R')
        pdf.ln(5)
        
        # Xác định kích thước trang và số cột
        page_width = 297 if orientation == 'L' else 210
        margin = 10
        usable_width = page_width - 2*margin
        
        # Chuyển đổi dữ liệu - loại bỏ Unicode/dấu tiếng Việt
        data = []
        headers = []
        
        # Xử lý tiêu đề cột - loại bỏ dấu
        for col in df.columns:
            # Chuyển thành ASCII, loại bỏ dấu tiếng Việt
            header_ascii = str(col).encode('ascii', 'ignore').decode('ascii')
            headers.append(header_ascii)
        
        # Xử lý dữ liệu - loại bỏ dấu tiếng Việt
        for _, row in df.iterrows():
            row_data = []
            for col in df.columns:
                # Chuyển thành chuỗi và loại bỏ dấu
                val = str(row[col]) if pd.notna(row[col]) else ""
                val_ascii = val.encode('ascii', 'ignore').decode('ascii')
                
                # Cắt bớt chuỗi dài
                if len(val_ascii) > 30:
                    val_ascii = val_ascii[:27] + "..."
                    
                row_data.append(val_ascii)
            data.append(row_data)
        
        # Tính toán độ rộng tối ưu cho mỗi cột
        col_widths = []
        
        # Font cho nội dung
        pdf.set_font('Arial', '', 8)
        
        for i, header in enumerate(headers):
            # Độ rộng tiêu đề
            header_width = pdf.get_string_width(header) + 6
            
            # Độ rộng nội dung
            max_content_width = 0
            for row in data:
                if i < len(row):
                    content_width = pdf.get_string_width(row[i]) + 6
                    max_content_width = max(max_content_width, content_width)
            
            # Lấy giá trị lớn nhất
            max_width = max(header_width, max_content_width)
            
            # Giới hạn độ rộng cột
            col_width = min(40, max(15, max_width))
            col_widths.append(col_width)
        
        # Điều chỉnh độ rộng cột để vừa với trang
        total_width = sum(col_widths)
        if total_width > usable_width:
            scale_factor = usable_width / total_width
            col_widths = [width * scale_factor for width in col_widths]
        
        # Vẽ tiêu đề cột
        pdf.set_font('Arial', 'B', 9)
        pdf.set_fill_color(240, 240, 240)
        
        for i, header in enumerate(headers):
            pdf.cell(col_widths[i], 10, header, 1, 0, 'C', 1)
        
        pdf.ln(10)
        
        # Vẽ dữ liệu
        pdf.set_font('Arial', '', 8)
        
        # Giới hạn số lượng hàng
        max_rows = min(1000, len(df))
        
        for i in range(max_rows):
            # Kiểm tra trang mới
            if pdf.get_y() + 7 > pdf.page_break_trigger:
                pdf.add_page()
                
                # Vẽ lại header sau khi chuyển trang
                pdf.set_font('Arial', 'B', 9)
                pdf.set_fill_color(240, 240, 240)
                
                for j, header in enumerate(headers):
                    pdf.cell(col_widths[j], 10, header, 1, 0, 'C', 1)
                pdf.ln(10)
                
                pdf.set_font('Arial', '', 8)
            
            # Vẽ dữ liệu hàng
            for j, width in enumerate(col_widths):
                if j < len(data[i]):
                    cell_text = data[i][j]
                    # Căn giữa cho số, căn trái cho text
                    align = 'C' if cell_text.replace('.', '', 1).isdigit() else 'L'
                    pdf.cell(width, 7, cell_text, 1, 0, align)
                else:
                    pdf.cell(width, 7, "", 1, 0, 'C')
            
            pdf.ln(7)
        
        # Thêm chân trang
        pdf.set_y(-20)
        pdf.set_font('Arial', 'I', 8)
        pdf.cell(0, 10, f'Trang {pdf.page_no()}/{"{nb}"}', 0, 0, 'C')
        pdf.cell(0, 10, 'He thong Khao sat & Danh gia', 0, 0, 'R')
        
        # Lưu PDF
        pdf.output(name=buffer, dest='S')
        buffer.seek(0)
        return buffer
    except Exception as e:
        print(f"Lỗi khi tạo PDF: {str(e)}")
        traceback.print_exc()
        
        # Tạo PDF lỗi đơn giản
        error_buffer = io.BytesIO()
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font('Arial', 'B', 16)
            pdf.cell(0, 10, 'Bao cao loi', 0, 1, 'C')
            pdf.set_font('Arial', '', 10)
            error_msg = f'Khong the tao PDF: {str(e)}'
            pdf.multi_cell(0, 10, error_msg, 0, 'L')
            pdf.output(name=error_buffer, dest='S')
            error_buffer.seek(0)
            return error_buffer
        except Exception as e2:
            print(f"Lỗi khi tạo báo cáo lỗi: {str(e2)}")
            empty_buffer = io.BytesIO()
            empty_buffer.write(b'%PDF-1.4\n%%EOF')  # Tạo PDF rỗng hợp lệ
            empty_buffer.seek(0)
            return empty_buffer

def create_student_report_docx(student_name, student_email, student_class, submission, questions, max_possible):
    """Tạo báo cáo chi tiết bài làm của học viên dạng DOCX"""
    try:
        doc = Document()
        
        # Thiết lập font chữ mặc định là Times New Roman
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Thêm tiêu đề - font chữ hỗ trợ Unicode
        heading = doc.add_heading(f"Báo cáo chi tiết - {student_name}", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thêm thời gian xuất báo cáo
        time_paragraph = doc.add_paragraph(f"Thời gian xuất báo cáo: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        time_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Thêm thông tin học viên
        doc.add_heading("Thông tin học viên", level=2)
        info_table = doc.add_table(rows=4, cols=2, style='Table Grid')
        
        # Đặt độ rộng cột
        for cell in info_table.columns[0].cells:
            cell.width = Inches(1.5)
        for cell in info_table.columns[1].cells:
            cell.width = Inches(4.5)
        
        # Thiết lập màu nền cho hàng tiêu đề
        for i in range(4):
            # Sửa lỗi: Đảm bảo có runs trước khi truy cập
            cell = info_table.rows[i].cells[0]
            cell_paragraph = cell.paragraphs[0]
            if not cell_paragraph.runs:
                cell_paragraph.add_run(cell.text if cell.text else '')
            cell_paragraph.runs[0].font.bold = True
            
            # Thêm màu nền
            shading_elm = parse_xml(r'<w:shd {} w:fill="E9E9E9"/>'.format(nsdecls('w')))
            info_table.rows[i].cells[0]._tc.get_or_add_tcPr().append(shading_elm)
        
        # Thêm dữ liệu vào bảng thông tin
        cells = info_table.rows[0].cells
        cells[0].text = "Họ và tên"
        cells[1].text = student_name
        
        cells = info_table.rows[1].cells
        cells[0].text = "Email"
        cells[1].text = student_email
        
        cells = info_table.rows[2].cells
        cells[0].text = "Lớp"
        cells[1].text = student_class
        
        # Xử lý timestamp tương thích với cả hai kiểu dữ liệu (số và chuỗi ISO)
        submission_time = "Không xác định"
        if isinstance(submission.get("timestamp"), (int, float)):
            try:
                submission_time = datetime.fromtimestamp(submission.get("timestamp")).strftime("%H:%M:%S %d/%m/%Y")
            except:
                pass
        else:
            try:
                dt = datetime.fromisoformat(submission.get("timestamp", "").replace("Z", "+00:00"))
                submission_time = dt.strftime("%H:%M:%S %d/%m/%Y")
            except:
                pass
        
        cells = info_table.rows[3].cells
        cells[0].text = "Thời gian nộp"
        cells[1].text = submission_time
        
        # Tính toán thông tin về bài làm
        total_correct = 0
        total_questions = len(questions)
        
        doc.add_heading("Chi tiết câu trả lời", level=2)
        
        # Tạo bảng chi tiết câu trả lời - cải thiện layout với cột rộng hợp lý
        answers_table = doc.add_table(rows=1, cols=5, style='Table Grid')
        
        # Thiết lập độ rộng tương đối cho các cột
        col_widths = [2.5, 2, 2, 1, 0.8]  # Tỷ lệ tương đối
        for i, width in enumerate(col_widths):
            for cell in answers_table.columns[i].cells:
                cell.width = Inches(width)
        
        # Thêm tiêu đề cho bảng với định dạng rõ ràng
        header_cells = answers_table.rows[0].cells
        headers = ["Câu hỏi", "Đáp án của học viên", "Đáp án đúng", "Kết quả", "Điểm"]
        
        # Tạo nền xám cho hàng tiêu đề
        for i, cell in enumerate(header_cells):
            cell.text = headers[i]
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Đảm bảo có runs trước khi truy cập
                if not paragraph.runs:
                    paragraph.add_run(headers[i])
                for run in paragraph.runs:
                    run.bold = True
            # Thêm màu nền
            shading_elm = parse_xml(r'<w:shd {} w:fill="E9E9E9"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Đảm bảo responses đúng định dạng
        responses = submission.get("responses", {})
        if isinstance(responses, str):
            try:
                responses = json.loads(responses)
            except:
                responses = {}
        
        # Thêm dữ liệu câu trả lời với định dạng cải thiện
        for q in questions:
            q_id = str(q.get("id", ""))
            
            # Đáp án người dùng
            user_ans = responses.get(q_id, [])
            
            # Chuẩn bị đáp án đúng
            q_correct = q.get("correct", [])
            q_answers = q.get("answers", [])
            
            if isinstance(q_correct, str):
                try:
                    q_correct = json.loads(q_correct)
                except:
                    try:
                        q_correct = [int(x.strip()) for x in q_correct.split(",")]
                    except:
                        q_correct = []
            
            if isinstance(q_answers, str):
                try:
                    q_answers = json.loads(q_answers)
                except:
                    q_answers = [q_answers]
            
            try:
                expected = [q_answers[i - 1] for i in q_correct]
            except (IndexError, TypeError):
                expected = ["Lỗi đáp án"]
            
            # Kiểm tra đúng/sai
            is_correct = check_answer_correctness(user_ans, q)
            if is_correct:
                total_correct += 1
                result = "Đúng"
                points = q.get("score", 0)
            else:
                result = "Sai"
                points = 0
            
            # Thêm hàng mới vào bảng
            row_cells = answers_table.add_row().cells
            
            # Thêm thông tin câu hỏi
            row_cells[0].text = f"Câu {q.get('id', '')}: {q.get('question', '')}"
            row_cells[1].text = ", ".join([str(a) for a in user_ans]) if user_ans else "Không trả lời"
            row_cells[2].text = ", ".join([str(a) for a in expected])
            row_cells[3].text = result
            
            # Đặt màu cho kết quả
            for paragraph in row_cells[3].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if not paragraph.runs:
                    paragraph.add_run(result)
                run = paragraph.runs[0]
                if is_correct:
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Màu xanh lá cho đúng
                    run.bold = True
                else:
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Màu đỏ cho sai
                    run.bold = True
            
            row_cells[4].text = str(points)
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thêm tổng kết với định dạng rõ ràng
        doc.add_heading("Tổng kết", level=2)
        summary_table = doc.add_table(rows=3, cols=2, style='Table Grid')
        
        # Thiết lập độ rộng cho bảng tổng kết
        for cell in summary_table.columns[0].cells:
            cell.width = Inches(1.5)
        for cell in summary_table.columns[1].cells:
            cell.width = Inches(3.0)
        
        # Thêm màu nền cho cột tiêu đề
        for i in range(3):
            cell = summary_table.rows[i].cells[0]
            paragraph = cell.paragraphs[0]
            if not paragraph.runs:
                paragraph.add_run(cell.text if cell.text else '')
            paragraph.runs[0].font.bold = True
            shading_elm = parse_xml(r'<w:shd {} w:fill="E9E9E9"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)
        
        cells = summary_table.rows[0].cells
        cells[0].text = "Số câu đúng"
        cells[1].text = f"{total_correct}/{total_questions}"
        
        cells = summary_table.rows[1].cells
        cells[0].text = "Điểm số"
        cells[1].text = f"{submission.get('score', 0)}/{max_possible}"
        
        cells = summary_table.rows[2].cells
        cells[0].text = "Tỷ lệ đúng"
        cells[1].text = f"{(total_correct/total_questions*100):.1f}%" if total_questions > 0 else "0%"
        
        # Thêm chân trang
        doc.add_paragraph()
        footer = doc.add_paragraph("Xuất báo cáo từ Hệ thống Khảo sát & Đánh giá")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        time_footer = doc.add_paragraph(f"Ngày xuất: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        time_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Lưu tệp
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    except Exception as e:
        print(f"Lỗi khi tạo báo cáo DOCX: {str(e)}")
        traceback.print_exc()
        # Trả về buffer trống nếu lỗi
        buffer = io.BytesIO()
        buffer.seek(0)
        return buffer

def create_student_report_pdf_fpdf(student_name, student_email, student_class, submission, questions, max_possible):
    """Tạo báo cáo chi tiết bài làm của học viên dạng PDF sử dụng FPDF với xử lý lỗi tiếng Việt"""
    buffer = io.BytesIO()
    
    try:
        # Loại bỏ dấu tiếng Việt từ tên
        student_name_ascii = student_name.encode('ascii', 'ignore').decode('ascii')
        student_class_ascii = student_class.encode('ascii', 'ignore').decode('ascii')
        
        # Tạo tiêu đề không dấu
        title = f"Bao cao chi tiet - {student_name_ascii}"
        
        # Tạo PDF mới 
        orientation = 'L' if len(questions) > 10 else 'P'
        pdf = FPDF(orientation=orientation, format='A4')
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        
        # Thêm tiêu đề
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(0, 10, title, 0, 1, 'C')
        
        # Thêm thời gian báo cáo
        pdf.set_font('Arial', 'I', 10)
        timestamp = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        pdf.cell(0, 5, f'Thoi gian xuat bao cao: {timestamp}', 0, 1, 'R')
        pdf.ln(5)
        
        # Tính toán thông tin về bài làm
        total_correct = 0
        total_questions = len(questions)
        
        # Đảm bảo responses đúng định dạng
        responses = submission.get("responses", {})
        if isinstance(responses, str):
            try:
                responses = json.loads(responses)
            except:
                responses = {}
        
        # Xử lý timestamp
        submission_time = "Không xác định"
        if isinstance(submission.get("timestamp"), (int, float)):
            try:
                submission_time = datetime.fromtimestamp(submission.get("timestamp")).strftime("%H:%M:%S %d/%m/%Y")
            except:
                pass
        else:
            try:
                dt = datetime.fromisoformat(submission.get("timestamp", "").replace("Z", "+00:00"))
                submission_time = dt.strftime("%H:%M:%S %d/%m/%Y")
            except:
                pass
        
        # Thông tin học viên - Tiêu đề
        pdf.set_font('Arial', 'B', 12)
        pdf.cell(0, 10, 'Thong tin hoc vien', 0, 1, 'L')
        
        # Bảng thông tin học viên
        pdf.set_font('Arial', '', 10)
        info_width = 190 if orientation == 'P' else 277
        col1_width = 50
        col2_width = info_width - col1_width
        
        # Tạo khung thông tin học viên
        pdf.set_fill_color(240, 240, 240)
        
        # Thông tin học viên - loại bỏ dấu tiếng Việt
        info_data = [
            ['Ho va ten', student_name_ascii],
            ['Email', student_email],
            ['Lop', student_class_ascii],
            ['Thoi gian nop', submission_time]
        ]
        
        for label, value in info_data:
            pdf.cell(col1_width, 10, label, 1, 0, 'L', 1)
            pdf.cell(col2_width, 10, value, 1, 1, 'L')
        
        pdf.ln(5)
        
        # Chi tiết câu trả lời - Tiêu đề
        pdf.set_font('Arial', 'B', 12)
        pdf.cell(0, 10, 'Chi tiet cau tra loi', 0, 1, 'L')
        
        # Tiêu đề bảng chi tiết
        pdf.set_font('Arial', 'B', 9)
        pdf.set_fill_color(240, 240, 240)
        
        # Xác định độ rộng cột - điều chỉnh phù hợp với nội dung và orientation
        if orientation == 'P':
            q_width = 70
            user_width = 35
            correct_width = 35
            result_width = 25
            points_width = 25
        else:
            q_width = 120
            user_width = 50
            correct_width = 50
            result_width = 30
            points_width = 27
        
        # Vẽ header bảng
        headers = ['Cau hoi', 'Dap an hoc vien', 'Dap an dung', 'Ket qua', 'Diem']
        widths = [q_width, user_width, correct_width, result_width, points_width]
        
        for i, header in enumerate(headers):
            pdf.cell(widths[i], 10, header, 1, 0, 'C', 1)
        pdf.ln(10)
        
        # Vẽ dữ liệu câu trả lời
        pdf.set_font('Arial', '', 9)
        
        for q in questions:
            q_id = str(q.get("id", ""))
            
            # Đáp án người dùng
            user_ans = responses.get(q_id, [])
            
            # Chuẩn bị đáp án đúng
            q_correct = q.get("correct", [])
            q_answers = q.get("answers", [])
            
            if isinstance(q_correct, str):
                try:
                    q_correct = json.loads(q_correct)
                except:
                    try:
                        q_correct = [int(x.strip()) for x in q_correct.split(",")]
                    except:
                        q_correct = []
            
            if isinstance(q_answers, str):
                try:
                    q_answers = json.loads(q_answers)
                except:
                    q_answers = [q_answers]
            
            try:
                expected = [q_answers[i - 1] for i in q_correct]
            except (IndexError, TypeError):
                expected = ["Loi dap an"]
            
            # Kiểm tra đúng/sai
            is_correct = check_answer_correctness(user_ans, q)
            if is_correct:
                total_correct += 1
                result = "Dung"
                points = q.get("score", 0)
            else:
                result = "Sai"
                points = 0
            
            # Chuẩn bị nội dung (loại bỏ dấu tiếng Việt)
            question_text = f"Cau {q.get('id', '')}: {q.get('question', '')}"
            question_text_ascii = question_text.encode('ascii', 'ignore').decode('ascii')
            
            # Giới hạn độ dài của các chuỗi
            if len(question_text_ascii) > (45 if orientation == 'P' else 80):
                question_text_ascii = question_text_ascii[:(42 if orientation == 'P' else 77)] + "..."
                
            user_answer_text = ", ".join([str(a) for a in user_ans]) if user_ans else "Khong tra loi"
            user_answer_text_ascii = user_answer_text.encode('ascii', 'ignore').decode('ascii')
            if len(user_answer_text_ascii) > (25 if orientation == 'P' else 40):
                user_answer_text_ascii = user_answer_text_ascii[:(22 if orientation == 'P' else 37)] + "..."
                
            correct_answer_text = ", ".join([str(a) for a in expected])
            correct_answer_text_ascii = correct_answer_text.encode('ascii', 'ignore').decode('ascii')
            if len(correct_answer_text_ascii) > (25 if orientation == 'P' else 40):
                correct_answer_text_ascii = correct_answer_text_ascii[:(22 if orientation == 'P' else 37)] + "..."
            
            # Kiểm tra phần còn lại của trang
            if pdf.get_y() + 10 > pdf.page_break_trigger:
                pdf.add_page()
                # Vẽ lại header sau khi chuyển trang
                pdf.set_font('Arial', 'B', 9)
                pdf.set_fill_color(240, 240, 240)
                for i, header in enumerate(headers):
                    pdf.cell(widths[i], 10, header, 1, 0, 'C', 1)
                pdf.ln(10)
                pdf.set_font('Arial', '', 9)
            
            # Vẽ dữ liệu
            pdf.cell(q_width, 10, question_text_ascii, 1, 0, 'L')
            pdf.cell(user_width, 10, user_answer_text_ascii, 1, 0, 'L')
            pdf.cell(correct_width, 10, correct_answer_text_ascii, 1, 0, 'L')
            
            # Đặt màu cho kết quả Đúng/Sai
            if is_correct:
                pdf.set_text_color(0, 128, 0)  # Màu xanh lá
            else:
                pdf.set_text_color(255, 0, 0)  # Màu đỏ
                
            pdf.cell(result_width, 10, result, 1, 0, 'C')
            
            # Đặt lại màu chữ cho điểm
            pdf.set_text_color(0, 0, 0)  # Màu đen
            pdf.cell(points_width, 10, str(points), 1, 1, 'C')
        
        pdf.ln(5)
        
        # Tổng kết
        pdf.set_font('Arial', 'B', 12)
        pdf.cell(0, 10, 'Tong ket', 0, 1, 'L')
        
        # Bảng tổng kết
        pdf.set_font('Arial', '', 10)
        pdf.set_fill_color(240, 240, 240)
        
        summary_col1 = 50
        summary_col2 = (190 if orientation == 'P' else 277) - summary_col1
        
        percent_correct = (total_correct/total_questions*100) if total_questions > 0 else 0
        summary_data = [
            ['So cau dung', f"{total_correct}/{total_questions}"],
            ['Diem so', f"{submission.get('score', 0)}/{max_possible}"],
            ['Ty le dung', f"{percent_correct:.1f}% {'(Dat)' if percent_correct >= 50 else '(Chua dat)'}"]
        ]
        
        for label, value in summary_data:
            pdf.cell(summary_col1, 10, label, 1, 0, 'L', 1)
            pdf.cell(summary_col2, 10, value, 1, 1, 'L')
        
        # Thêm chân trang
        pdf.set_y(-20)
        pdf.set_font('Arial', 'I', 8)
        pdf.cell(0, 10, f'Trang {pdf.page_no()}/{"{nb}"}', 0, 0, 'C')
        pdf.cell(0, 10, 'He thong Khao sat & Danh gia', 0, 0, 'R')
        
        # Lưu PDF vào buffer
        pdf.output(name=buffer, dest='S')
        
    except Exception as e:
        print(f"Lỗi khi tạo báo cáo PDF: {str(e)}")
        traceback.print_exc()
        
        # Tạo báo cáo đơn giản nếu gặp lỗi
        try:
            error_buffer = io.BytesIO()
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font('Arial', 'B', 16)
            pdf.cell(0, 10, f'Bao cao chi tiet - {student_name.encode("ascii", "ignore").decode("ascii")}', 0, 1, 'C')
            pdf.set_font('Arial', '', 10)
            error_text = f'Khong the hien thi bao cao chi tiet voi font tieng Viet.\nSu dung dinh dang DOCX de xem day du.\nLoi: {str(e)}'
            pdf.multi_cell(0, 10, error_text, 0, 'L')
            pdf.output(name=error_buffer, dest='S')
            error_buffer.seek(0)
            return error_buffer
        except Exception as e2:
            print(f"Không thể tạo báo cáo thay thế: {str(e2)}")
            empty_buffer = io.BytesIO()
            empty_buffer.write(b'%PDF-1.4\n%%EOF')  # Tạo PDF rỗng hợp lệ
            empty_buffer.seek(0)
            return empty_buffer
    
    buffer.seek(0)
    return buffer

def display_overview_tab(submissions=None, students=None, questions=None, max_possible=0):
    """Hiển thị tab tổng quan"""
    if submissions is None:
        submissions = []
    if students is None:
        students = []
    if questions is None:
        questions = []
        
    st.subheader("Tổng quan kết quả")
    
    # Thống kê cơ bản
    total_submissions = len(submissions)
    if total_submissions > 0:
        scores = [s.get("score", 0) for s in submissions]
        avg_score = sum(scores) / total_submissions
        max_score = max(scores) if scores else 0
    else:
        avg_score = 0
        max_score = 0
        
    total_users = len(set([s.get("user_email") for s in submissions]))
    
    # Hiển thị metrics
    col1, col2, col3 = st.columns(3)
    col1.metric("📝 Tổng số bài nộp", total_submissions)
    col1.metric("👥 Số học viên đã làm", total_users)
    
    col2.metric("📊 Điểm trung bình", f"{avg_score:.2f}/{max_possible}")
    col2.metric("🏆 Điểm cao nhất", f"{max_score}/{max_possible}")
    
    col3.metric("📋 Số câu hỏi", len(questions))
    col3.metric("👨‍🎓 Tổng số học viên", len(students))
    
    # Biểu đồ điểm số theo thời gian
    st.subheader("Điểm số theo thời gian")
    
    # Chuẩn bị dữ liệu
    time_data = []
    for s in submissions:
        # Xử lý timestamp
        try:
            if isinstance(s.get("timestamp"), (int, float)):
                submit_time = datetime.fromtimestamp(s.get("timestamp"))
            else:
                submit_time = datetime.fromisoformat(s.get("timestamp", "").replace("Z", "+00:00"))
            
            time_data.append({
                "timestamp": submit_time,
                "score": s.get("score", 0),
                "user": s.get("user_email", "")
            })
        except:
            # Bỏ qua bài nộp có timestamp không hợp lệ
            pass
    
    if time_data:
        df_time = pd.DataFrame(time_data)
        if not df_time.empty:
            df_time = df_time.sort_values("timestamp")
            
            # Vẽ biểu đồ
            fig, ax = plt.subplots(figsize=(10, 5))
            ax.plot(df_time["timestamp"], df_time["score"], marker='o')
            ax.set_ylabel("Điểm số")
            ax.set_xlabel("Thời gian nộp bài")
            ax.grid(True, linestyle='--', alpha=0.7)
            
            # Giảm số lượng tick trên trục x
            max_ticks = 6
            if len(df_time) > max_ticks:
                stride = len(df_time) // max_ticks
                plt.xticks(df_time["timestamp"][::stride])
            
            # Sử dụng constrained_layout thay vì tight_layout
            fig.set_constrained_layout(True)
            st.pyplot(fig)
            
            # Hiển thị phân phối điểm
            st.subheader("Phân phối điểm số")
            if submissions:
                scores = [s.get("score", 0) for s in submissions]
                fig, ax = plt.subplots(figsize=(10, 5))
                ax.hist(scores, bins=min(10, len(set(scores))), alpha=0.7, color='skyblue', edgecolor='black')
                ax.set_xlabel("Điểm số")
                ax.set_ylabel("Số lượng bài nộp")
                ax.grid(True, linestyle='--', alpha=0.3)
                fig.set_constrained_layout(True)
                st.pyplot(fig)
    else:
        st.info("Không có đủ dữ liệu để vẽ biểu đồ theo thời gian.")

def display_student_tab(submissions=None, students=None, questions=None, max_possible=0):
    """Hiển thị tab theo học viên"""
    if submissions is None:
        submissions = []
    if students is None:
        students = []
    if questions is None:
        questions = []
        
    st.subheader("Chi tiết theo học viên")
    
    # Tạo DataFrame từ dữ liệu
    user_data = []
    for s in submissions:
        try:
            # Tìm thông tin học viên
            student_info = next((student for student in students if student.get("email") == s.get("user_email")), None)
            full_name = student_info.get("full_name", "Không xác định") if student_info else "Không xác định"
            class_name = student_info.get("class", "Không xác định") if student_info else "Không xác định"
            
            # Xử lý timestamp
            submission_time = "Không xác định"
            try:
                if isinstance(s.get("timestamp"), (int, float)):
                    submission_time = datetime.fromtimestamp(s.get("timestamp")).strftime("%H:%M:%S %d/%m/%Y")
                else:
                    dt = datetime.fromisoformat(s.get("timestamp", "").replace("Z", "+00:00"))
                    submission_time = dt.strftime("%H:%M:%S %d/%m/%Y")
            except:
                pass
            
            user_data.append({
                "email": s.get("user_email", ""),
                "full_name": full_name,
                "class": class_name,
                "submission_id": s.get("id", ""),
                "timestamp": submission_time,
                "score": s.get("score", 0),
                "max_score": max_possible,
                "percent": f"{(s.get('score', 0)/max_possible*100):.1f}%" if max_possible > 0 else "N/A"
            })
        except Exception as e:
            st.error(f"Lỗi khi xử lý dữ liệu học viên: {str(e)}")
    
    if user_data:
        df_users = pd.DataFrame(user_data)
        
        # Lọc theo email hoặc lớp
        col1, col2 = st.columns(2)
        with col1:
            user_filter = st.selectbox(
                "Chọn học viên để xem chi tiết:",
                options=["Tất cả"] + sorted(list(set([u.get("email", "") for u in user_data]))),
                key="user_filter_tab2"
            )
        
        with col2:
            unique_classes = [u.get("class", "") for u in user_data if u.get("class") != "Không xác định"]
            class_filter = st.selectbox(
                "Lọc theo lớp:",
                options=["Tất cả"] + sorted(list(set(unique_classes))),
                key="class_filter_tab2"
            )
        
        # Áp dụng bộ lọc
        df_filtered = df_users
        
        if user_filter != "Tất cả":
            df_filtered = df_filtered[df_filtered["email"] == user_filter]
        
        if class_filter != "Tất cả":
            df_filtered = df_filtered[df_filtered["class"] == class_filter]
        
        # Hiển thị bảng
        st.dataframe(
            df_filtered.sort_values(by="timestamp", ascending=False),
            use_container_width=True,
            hide_index=True
        )
        
        # Xem chi tiết một bài nộp cụ thể
        if user_filter != "Tất cả":
            submission_ids = df_filtered["submission_id"].tolist()
            if submission_ids:
                selected_submission = st.selectbox(
                    "Chọn bài nộp để xem chi tiết:",
                    options=submission_ids,
                    key="submission_id_select"
                )
                
                # Tìm bài nộp được chọn
                submission = next((s for s in submissions if str(s.get("id", "")) == str(selected_submission)), None)
                if submission:
                    st.subheader(f"Chi tiết bài nộp #{selected_submission}")
                    
                    total_correct = 0
                    total_questions = len(questions)
                    student_detail_data = []
                    
                    # Đảm bảo responses đúng định dạng
                    responses = submission.get("responses", {})
                    if isinstance(responses, str):
                        try:
                            responses = json.loads(responses)
                        except:
                            responses = {}
                    
                    # Hiển thị câu trả lời chi tiết
                    for q in questions:
                        q_id = str(q.get("id", ""))
                        st.write(f"**Câu {q.get('id', '')}: {q.get('question', '')}**")
                        
                        # Đáp án người dùng
                        user_ans = responses.get(q_id, [])
                        
                        # Chuẩn bị dữ liệu đáp án đúng
                        q_correct = q.get("correct", [])
                        q_answers = q.get("answers", [])
                        
                        if isinstance(q_correct, str):
                            try:
                                q_correct = json.loads(q_correct)
                            except:
                                try:
                                    q_correct = [int(x.strip()) for x in q_correct.split(",")]
                                except:
                                    q_correct = []
                        
                        if isinstance(q_answers, str):
                            try:
                                q_answers = json.loads(q_answers)
                            except:
                                q_answers = [q_answers]
                        
                        try:
                            expected = [q_answers[i - 1] for i in q_correct]
                        except (IndexError, TypeError):
                            expected = ["Lỗi đáp án"]
                        
                        # Kiểm tra đúng/sai
                        is_correct = check_answer_correctness(user_ans, q)
                        if is_correct:
                            total_correct += 1
                        
                        # Thu thập dữ liệu chi tiết
                        student_detail_data.append({
                            "Câu hỏi": f"Câu {q.get('id', '')}: {q.get('question', '')}",
                            "Đáp án của học viên": ", ".join([str(a) for a in user_ans]) if user_ans else "Không trả lời",
                            "Đáp án đúng": ", ".join([str(a) for a in expected]),
                            "Kết quả": "Đúng" if is_correct else "Sai",
                            "Điểm": q.get("score", 0) if is_correct else 0
                        })
                        
                        # Hiển thị đáp án của người dùng
                        st.write("Đáp án của học viên:")
                        if not user_ans:
                            st.write("- Không trả lời")
                        else:
                            for ans in user_ans:
                                st.write(f"- {ans}")
                        
                        # Hiển thị kết quả
                        if is_correct:
                            st.success(f"✅ Đúng (+{q.get('score', 0)} điểm)")
                        else:
                            st.error("❌ Sai (0 điểm)")
                            st.write("Đáp án đúng:")
                            for ans in expected:
                                st.write(f"- {ans}")
                        
                        st.divider()
                    
                    # Hiển thị thống kê tổng hợp
                    st.subheader("Tổng kết")
                    col1, col2, col3 = st.columns(3)
                    col1.metric("Số câu đúng", f"{total_correct}/{total_questions}")
                    col2.metric("Điểm số", f"{submission.get('score', 0)}/{max_possible}")
                    col3.metric("Tỷ lệ đúng", f"{(total_correct/total_questions*100):.1f}%" if total_questions > 0 else "0%")
                    
                    # Tạo DataFrame chi tiết
                    df_student_detail = pd.DataFrame(student_detail_data)
                    
                    # Xuất báo cáo chi tiết
                    st.write("### Xuất báo cáo chi tiết")
                    
                    # Người dùng và thông tin
                    student_info = next((student for student in students if student.get("email") == submission.get("user_email")), None)
                    student_name = student_info.get("full_name", "Không xác định") if student_info else "Không xác định"
                    student_class = student_info.get("class", "Không xác định") if student_info else "Không xác định"
                    
                    # Tạo báo cáo chi tiết
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        # Tạo báo cáo dạng DOCX
                        try:
                            docx_buffer = create_student_report_docx(
                                student_name,
                                submission.get("user_email", ""),
                                student_class,
                                submission,
                                questions,
                                max_possible
                            )
                            
                            st.markdown(
                                get_download_link_docx(docx_buffer, 
                                                    f"bao_cao_{student_name.replace(' ', '_')}_{submission.get('id', '')}.docx", 
                                                    "Tải xuống báo cáo chi tiết (DOCX)"), 
                                unsafe_allow_html=True
                            )
                        except Exception as e:
                            st.error(f"Không thể tạo báo cáo DOCX: {str(e)}")
                    
                    with col2:
                        # Tạo báo cáo dạng PDF
                        try:
                            pdf_buffer = create_student_report_pdf_fpdf(
                                student_name,
                                submission.get("user_email", ""),
                                student_class,
                                submission,
                                questions,
                                max_possible
                            )
                            
                            st.markdown(
                                get_download_link_pdf(pdf_buffer, 
                                                    f"bao_cao_{student_name.replace(' ', '_')}_{submission.get('id', '')}.pdf", 
                                                    "Tải xuống báo cáo chi tiết (PDF)"), 
                                unsafe_allow_html=True
                            )
                        except Exception as e:
                            st.error(f"Không thể tạo báo cáo PDF: {str(e)}")
    else:
        st.info("Không có dữ liệu học viên để hiển thị.")

def display_question_tab(submissions=None, questions=None):
    """Hiển thị tab phân tích câu hỏi"""
    if submissions is None:
        submissions = []
    if questions is None:
        questions = []
        
    st.subheader("Phân tích theo câu hỏi")
    
    # Thống kê tỷ lệ đúng/sai cho từng câu hỏi
    question_stats = {}
    
    for q in questions:
        q_id = str(q.get("id", ""))
        correct_count = 0
        wrong_count = 0
        skip_count = 0
        
        for s in submissions:
            # Đảm bảo responses đúng định dạng
            responses = s.get("responses", {})
            if isinstance(responses, str):
                try:
                    responses = json.loads(responses)
                except:
                    responses = {}
            
            user_ans = responses.get(q_id, [])
            
            if not user_ans:
                skip_count += 1
            elif check_answer_correctness(user_ans, q):
                correct_count += 1
            else:
                wrong_count += 1
        
        question_stats[q_id] = {
            "question": q.get("question", ""),
            "correct": correct_count,
            "wrong": wrong_count,
            "skip": skip_count,
            "total": correct_count + wrong_count + skip_count,
            "correct_rate": correct_count / (correct_count + wrong_count + skip_count) if (correct_count + wrong_count + skip_count) > 0 else 0
        }
    
    # DataFrame thống kê câu hỏi
    df_questions_data = [
        {
            "Câu hỏi ID": q_id,
            "Nội dung": stats["question"],
            "Số lượng đúng": stats["correct"],
            "Số lượng sai": stats["wrong"],
            "Bỏ qua": stats["skip"],
            "Tổng số làm": stats["total"],
            "Tỷ lệ đúng (%)": f"{stats['correct_rate']*100:.1f}%"
        }
        for q_id, stats in question_stats.items()
    ]
    
    if not df_questions_data:
        st.info("Không có dữ liệu câu hỏi để phân tích.")
        return pd.DataFrame()
    
    df_questions = pd.DataFrame(df_questions_data)
    
    # Vẽ biểu đồ tỷ lệ đúng theo từng câu hỏi
    q_ids = list(question_stats.keys())
    correct_rates = [question_stats[q_id]["correct_rate"] * 100 for q_id in q_ids]
    
    # Giới hạn độ dài câu hỏi để hiển thị trên biểu đồ
    short_questions = [f"Câu {q_id}: {question_stats[q_id]['question'][:30]}..." for q_id in q_ids]
    
    # Tạo biểu đồ với kích thước nhỏ hơn
    fig, ax = plt.subplots(figsize=(10, 6))
    bars = ax.bar(short_questions, correct_rates, color='skyblue')
    
    # Xoay nhãn để tránh chồng chéo
    plt.xticks(rotation=45, ha='right')
    
    # Thêm nhãn giá trị
    for bar in bars:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height + 1,
                f'{height:.1f}%', ha='center', va='bottom', fontsize=9)
    
    ax.set_ylim(0, 105)  # Giới hạn trục y từ 0-100%
    ax.set_xlabel("Câu hỏi")
    ax.set_ylabel("Tỷ lệ đúng (%)")
    ax.set_title("Tỷ lệ trả lời đúng theo từng câu hỏi")
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    fig.set_constrained_layout(True)
    st.pyplot(fig)
    
    # Hiển thị bảng thống kê
    st.dataframe(df_questions, use_container_width=True, hide_index=True)
    
    # Chi tiết từng câu hỏi
    if q_ids:
        selected_question = st.selectbox(
            "Chọn câu hỏi để xem chi tiết:",
            options=[(f"Câu {q_id}: {question_stats[q_id]['question']}") for q_id in q_ids],
            key="question_select_tab3"
        )
        
        if selected_question:
            try:
                q_id = selected_question.split(":")[0].replace("Câu ", "").strip()
                q_data = question_stats[q_id]
                q_detail = next((q for q in questions if str(q.get("id", "")) == q_id), None)
                
                if q_detail:
                    st.write(f"**{selected_question}**")
                    
                    # Hiển thị thống kê
                    col1, col2, col3, col4 = st.columns(4)
                    col1.metric("✅ Đúng", q_data["correct"])
                    col2.metric("❌ Sai", q_data["wrong"])
                    col3.metric("⏭️ Bỏ qua", q_data["skip"])
                    col4.metric("📊 Tỷ lệ đúng", f"{q_data['correct_rate']*100:.1f}%")
                    
                    # Tạo biểu đồ tròn nhỏ hơn
                    fig, ax = plt.subplots(figsize=(6, 4))
                    
                    # Sử dụng biểu đồ đơn giản với nhãn và tỷ lệ bên ngoài
                    labels = ['Đúng', 'Sai', 'Bỏ qua']
                    sizes = [q_data["correct"], q_data["wrong"], q_data["skip"]]
                    colors = ['#4CAF50', '#F44336', '#9E9E9E']
                    
                    # Chỉ hiển thị phần trăm nếu giá trị > 0
                    patches, texts, autotexts = ax.pie(
                        sizes, 
                        labels=None,  # Không hiển thị nhãn trên biểu đồ
                        colors=colors, 
                        autopct=lambda p: f'{p:.1f}%' if p > 0 else '',
                        startangle=90,
                        pctdistance=0.85  # Đặt phần trăm gần hơn với trung tâm
                    )
                    
                    # Thiết lập kích thước font nhỏ hơn
                    for autotext in autotexts:
                        autotext.set_fontsize(9)
                    
                    # Thêm chú thích bên ngoài biểu đồ
                    ax.legend(labels, loc="upper right", fontsize=9)
                    
                    # Vẽ vòng tròn trắng ở giữa
                    centre_circle = plt.Circle((0, 0), 0.5, fc='white')
                    ax.add_patch(centre_circle)
                    
                    ax.axis('equal')  # Giữ tỷ lệ vòng tròn
                    fig.set_constrained_layout(True)
                    st.pyplot(fig)
                    
                    # Đáp án đúng
                    st.write("**Đáp án đúng:**")
                    
                    # Chuẩn bị dữ liệu đáp án đúng
                    q_correct = q_detail.get("correct", [])
                    q_answers = q_detail.get("answers", [])
                    
                    if isinstance(q_correct, str):
                        try:
                            q_correct = json.loads(q_correct)
                        except:
                            try:
                                q_correct = [int(x.strip()) for x in q_correct.split(",")]
                            except:
                                q_correct = []
                    
                    if isinstance(q_answers, str):
                        try:
                            q_answers = json.loads(q_answers)
                        except:
                            q_answers = [q_answers]
                    
                    try:
                        for i in q_correct:
                            st.write(f"- {q_answers[i-1]}")
                    except (IndexError, TypeError):
                        st.write("- Lỗi hiển thị đáp án")
            except Exception as e:
                st.error(f"Lỗi khi hiển thị chi tiết câu hỏi: {str(e)}")
    
    return df_questions

def display_student_list_tab(submissions=None, students=None, max_possible=0):
    """Hiển thị tab danh sách học viên"""
    if submissions is None:
        submissions = []
    if students is None:
        students = []
        
    st.subheader("Danh sách học viên")
    
    if not students:
        st.info("Chưa có học viên nào đăng ký")
        return pd.DataFrame(), pd.DataFrame()
    
    # Chuẩn bị dữ liệu
    student_data = []
    for student in students:
        try:
            # Tìm tất cả bài nộp của học viên
            student_email = student.get("email", "")
            student_submissions = [s for s in submissions if s.get("user_email") == student_email]
            submission_count = len(student_submissions)
            
            # Tìm điểm cao nhất
            max_student_score = max([s.get("score", 0) for s in student_submissions]) if student_submissions else 0
            
            # Thời gian đăng ký
            registration_date = format_date(student.get("registration_date"))
            
            student_data.append({
                "full_name": student.get("full_name", ""),
                "email": student_email,
                "class": student.get("class", ""),
                "registration_date": registration_date,
                "submission_count": submission_count,
                "max_score": max_student_score,
                "max_possible": max_possible,
                "percent": f"{(max_student_score/max_possible*100):.1f}%" if max_possible > 0 else "N/A"
            })
        except Exception as e:
            st.error(f"Lỗi khi xử lý dữ liệu học viên {student.get('email', '')}: {str(e)}")
    
    # DataFrame cho danh sách học viên
    students_list_data = [
        {
            "Họ và tên": s["full_name"],
            "Email": s["email"],
            "Lớp": s["class"],
            "Ngày đăng ký": s["registration_date"],
            "Số lần làm bài": s["submission_count"],
            "Điểm cao nhất": s["max_score"],
            "Điểm tối đa": s["max_possible"],
            "Tỷ lệ đúng": s["percent"]
        } for s in student_data
    ]
    
    df_students_list = pd.DataFrame(students_list_data)
    
    # Lọc theo lớp
    class_options = ["Tất cả"] + sorted(list(set([s["class"] for s in student_data if s["class"]])))
    class_filter = st.selectbox(
        "Lọc theo lớp:",
        options=class_options,
        key="class_filter_tab4"
    )
    
    df_students = pd.DataFrame(student_data)
    
    if class_filter != "Tất cả":
        df_students = df_students[df_students["class"] == class_filter]
    
    # Sắp xếp theo tên
    df_students = df_students.sort_values(by="full_name")
    
    # Hiển thị bảng
    st.dataframe(
        df_students,
        use_container_width=True,
        hide_index=True
    )
    
    # Thống kê theo lớp
    st.subheader("Thống kê theo lớp")
    
    # Nhóm theo lớp
    df_class_stats = pd.DataFrame()
    if not df_students.empty and "class" in df_students.columns:
        # Đảm bảo rằng class không rỗng
        df_students["class"] = df_students["class"].fillna("Không xác định")
        
        class_stats = df_students.groupby("class").agg({
            "email": "count",
            "submission_count": "sum",
            "max_score": "mean"
        }).reset_index()
        
        class_stats.columns = ["Lớp", "Số học viên", "Tổng số bài nộp", "Điểm trung bình"]
        class_stats["Điểm trung bình"] = class_stats["Điểm trung bình"].round(2)
        
        # DataFrame thống kê lớp
        df_class_stats = class_stats.copy()
        
        st.dataframe(
            class_stats,
            use_container_width=True,
            hide_index=True
        )
        
        # Biểu đồ cột nhỏ hơn cho số học viên theo lớp
        fig, ax = plt.subplots(figsize=(10, 4))
        ax.bar(class_stats["Lớp"], class_stats["Số học viên"], color='skyblue')
        ax.set_xlabel("Lớp")
        ax.set_ylabel("Số học viên")
        ax.set_title("Số học viên theo lớp")
        plt.xticks(rotation=45, ha='right')
        fig.set_constrained_layout(True)
        st.pyplot(fig)
    else:
        st.info("Không có đủ dữ liệu để hiển thị thống kê theo lớp.")
    
    return df_students_list, df_class_stats

def display_export_tab(df_all_submissions=None, df_questions=None, df_students_list=None, df_class_stats=None):
    """Hiển thị tab xuất báo cáo"""
    if df_all_submissions is None:
        df_all_submissions = pd.DataFrame()
    if df_questions is None:
        df_questions = pd.DataFrame()
    if df_students_list is None:
        df_students_list = pd.DataFrame()
    if df_class_stats is None:
        df_class_stats = pd.DataFrame()
        
    st.subheader("Xuất báo cáo")
    
    # Kiểm tra phiên bản FPDF
    try:
        import pkg_resources
        fpdf_version = pkg_resources.get_distribution("fpdf").version
        is_fpdf2 = fpdf_version.startswith("2.")
        
        if not is_fpdf2:
            st.warning("""
            ⚠️ Lưu ý: Bạn đang sử dụng FPDF1 không hỗ trợ tiếng Việt có dấu. 
            Các báo cáo PDF sẽ không hiển thị được tiếng Việt có dấu hoặc có thể gặp lỗi khi mở.
            Khuyên dùng định dạng DOCX hoặc cài đặt fpdf2: `pip install fpdf2`
            """)
    except:
        st.warning("Không xác định được phiên bản FPDF. Có thể gặp lỗi khi xuất PDF.")
    
     
    # Thêm tab cho các loại báo cáo khác nhau
    report_tab1, report_tab2 = st.tabs(["Báo cáo tổng hợp", "Báo cáo theo học viên"])
    
    with report_tab1:
        # Hiển thị các loại báo cáo có thể xuất
        if not df_all_submissions.empty:
            st.write("### 1. Báo cáo tất cả bài nộp")
            
            col1, col2 = st.columns(2)
            
            with col1:
                try:
                    # DOCX
                    docx_buffer = dataframe_to_docx(df_all_submissions, "Báo cáo tất cả bài nộp", "bao_cao_tat_ca_bai_nop.docx")
                    st.markdown(get_download_link_docx(docx_buffer, "bao_cao_tat_ca_bai_nop.docx", 
                                                "Tải xuống báo cáo (DOCX)"), unsafe_allow_html=True)
                except Exception as e:
                    st.error(f"Lỗi khi tạo DOCX: {str(e)}")
            
            with col2:
                try:
                    # PDF - sử dụng FPDF thay vì ReportLab
                    pdf_buffer = dataframe_to_pdf_fpdf(df_all_submissions, "Báo cáo tất cả bài nộp", "bao_cao_tat_ca_bai_nop.pdf")
                    st.markdown(get_download_link_pdf(pdf_buffer, "bao_cao_tat_ca_bai_nop.pdf", 
                                                "Tải xuống báo cáo (PDF)"), unsafe_allow_html=True)
                except Exception as e:
                    st.error(f"Lỗi khi tạo PDF: {str(e)}")
        
        # ... [tiếp tục cho các phần báo cáo khác]
        if not df_questions.empty:
            st.write("### 2. Báo cáo thống kê câu hỏi")
            
            col1, col2 = st.columns(2)
            
            with col1:
                try:
                    # DOCX
                    docx_buffer = dataframe_to_docx(df_questions, "Báo cáo thống kê câu hỏi", "bao_cao_thong_ke_cau_hoi.docx")
                    st.markdown(get_download_link_docx(docx_buffer, "bao_cao_thong_ke_cau_hoi.docx", 
                                                "Tải xuống báo cáo (DOCX)"), unsafe_allow_html=True)
                except Exception as e:
                    st.error(f"Lỗi khi tạo DOCX: {str(e)}")
            
            with col2:
                try:
                    # PDF
                    pdf_buffer = dataframe_to_pdf_fpdf(df_questions, "Báo cáo thống kê câu hỏi", "bao_cao_thong_ke_cau_hoi.pdf")
                    st.markdown(get_download_link_pdf(pdf_buffer, "bao_cao_thong_ke_cau_hoi.pdf", 
                                                "Tải xuống báo cáo (PDF)"), unsafe_allow_html=True)
                except Exception as e:
                    st.error(f"Lỗi khi tạo PDF: {str(e)}")
        
        if not df_students_list.empty:
            st.write("### 3. Báo cáo danh sách học viên")
            
            col1, col2 = st.columns(2)
            
            with col1:
                try:
                    # DOCX
                    docx_buffer = dataframe_to_docx(df_students_list, "Báo cáo danh sách học viên", "bao_cao_danh_sach_hoc_vien.docx")
                    st.markdown(get_download_link_docx(docx_buffer, "bao_cao_danh_sach_hoc_vien.docx", 
                                                "Tải xuống báo cáo (DOCX)"), unsafe_allow_html=True)
                except Exception as e:
                    st.error(f"Lỗi khi tạo DOCX: {str(e)}")
            
            with col2:
                try:
                    # PDF
                    pdf_buffer = dataframe_to_pdf_fpdf(df_students_list, "Báo cáo danh sách học viên", "bao_cao_danh_sach_hoc_vien.pdf")
                    st.markdown(get_download_link_pdf(pdf_buffer, "bao_cao_danh_sach_hoc_vien.pdf", 
                                                "Tải xuống báo cáo (PDF)"), unsafe_allow_html=True)
                except Exception as e:
                    st.error(f"Lỗi khi tạo PDF: {str(e)}")
        
        if not df_class_stats.empty:
            st.write("### 4. Báo cáo thống kê theo lớp")
            
            col1, col2 = st.columns(2)
            
            with col1:
                try:
                    # DOCX
                    docx_buffer = dataframe_to_docx(df_class_stats, "Báo cáo thống kê theo lớp", "bao_cao_thong_ke_lop.docx")
                    st.markdown(get_download_link_docx(docx_buffer, "bao_cao_thong_ke_lop.docx", 
                                                "Tải xuống báo cáo (DOCX)"), unsafe_allow_html=True)
                except Exception as e:
                    st.error(f"Lỗi khi tạo DOCX: {str(e)}")
            
            with col2:
                try:
                    # PDF
                    pdf_buffer = dataframe_to_pdf_fpdf(df_class_stats, "Báo cáo thống kê theo lớp", "bao_cao_thong_ke_lop.pdf")
                    st.markdown(get_download_link_pdf(pdf_buffer, "bao_cao_thong_ke_lop.pdf", 
                                                "Tải xuống báo cáo (PDF)"), unsafe_allow_html=True)
                except Exception as e:
                    st.error(f"Lỗi khi tạo PDF: {str(e)}")
        
        st.write("### 5. Báo cáo tổng hợp (Excel)")
        
        try:
            # Chuẩn bị danh sách DataFrame và tên sheet
            dfs = []
            sheet_names = []
            
            if not df_all_submissions.empty:
                dfs.append(df_all_submissions)
                sheet_names.append("Tất cả bài nộp")
            
            if not df_questions.empty:
                dfs.append(df_questions)
                sheet_names.append("Thống kê câu hỏi")
            
            if not df_students_list.empty:
                dfs.append(df_students_list)
                sheet_names.append("Danh sách học viên")
            
            if not df_class_stats.empty:
                dfs.append(df_class_stats)
                sheet_names.append("Thống kê lớp")
            
            if dfs and sheet_names:
                # Hiển thị link tải xuống
                st.markdown(export_to_excel(dfs, sheet_names, "bao_cao_tong_hop.xlsx"), unsafe_allow_html=True)
            else:
                st.info("Không có đủ dữ liệu để tạo báo cáo Excel.")
            
        except Exception as e:
            st.error(f"Lỗi khi tạo file Excel: {str(e)}")

    
    with report_tab2:
        st.write("### Báo cáo chi tiết theo từng học viên")
        
        # Lấy danh sách học viên và bài nộp từ database
        try:
            # Lấy dữ liệu từ database
            supabase = get_supabase_client()
            if not supabase:
                st.error("Không thể kết nối đến Supabase.")
                return
                
            students = get_all_users(role="Học viên")
            questions = get_all_questions()
            max_possible = sum([q.get("score", 0) for q in questions])
            
            # Lấy danh sách email học viên từ các bài nộp
            student_emails = []
            for student in students:
                student_email = student.get("email", "")
                if student_email:
                    student_emails.append(student_email)
            
            student_emails = sorted(student_emails)
                
            if not student_emails:
                st.info("Không có dữ liệu học viên để hiển thị.")
                return
            
            # Hiển thị dropdown để chọn học viên
            selected_email = st.selectbox(
                "Chọn email học viên:",
                options=student_emails
            )
            
            if selected_email:
                # Lấy thông tin học viên
                student_info = next((student for student in students if student.get("email") == selected_email), None)
                if not student_info:
                    st.warning(f"Không tìm thấy thông tin học viên: {selected_email}")
                    return
                    
                student_name = student_info.get("full_name", "Không xác định")
                student_class = student_info.get("class", "Không xác định")
                
                # Lấy tất cả bài nộp của học viên này
                student_submissions = get_user_submissions(selected_email)
                
                if student_submissions:
                    st.success(f"Đã tìm thấy {len(student_submissions)} bài làm của học viên {student_name} ({selected_email})")
                    
                    # Hiển thị thông tin tổng quan
                    max_score = max([s.get("score", 0) for s in student_submissions]) if student_submissions else 0
                    best_submission = max(student_submissions, key=lambda x: x.get("score", 0))
                    best_score = best_submission.get("score", 0)
                    best_percent = (best_score / max_possible) * 100 if max_possible > 0 else 0
                    
                    col1, col2 = st.columns(2)
                    col1.metric("Số lần làm bài", len(student_submissions))
                    col2.metric("Điểm cao nhất", f"{best_score}/{max_possible} ({best_percent:.1f}%)")
                    
                    # Tạo DataFrame cho xuất báo cáo
                    student_report_data = []
                    
                    for idx, submission in enumerate(student_submissions):
                        # Xử lý timestamp
                        submission_time = "Không xác định"
                        if isinstance(submission.get("timestamp"), (int, float)):
                            try:
                                submission_time = datetime.fromtimestamp(submission.get("timestamp")).strftime("%H:%M:%S %d/%m/%Y")
                            except:
                                pass
                        else:
                            try:
                                dt = datetime.fromisoformat(submission.get("timestamp", "").replace("Z", "+00:00"))
                                submission_time = dt.strftime("%H:%M:%S %d/%m/%Y")
                            except:
                                pass
                        
                        # Đảm bảo responses đúng định dạng
                        responses = submission.get("responses", {})
                        if isinstance(responses, str):
                            try:
                                responses = json.loads(responses)
                            except:
                                responses = {}
                        
                        # Tính số câu trả lời đúng
                        correct_count = 0
                        for q in questions:
                            q_id = str(q.get("id", ""))
                            user_ans = responses.get(q_id, [])
                            is_correct = check_answer_correctness(user_ans, q)
                            if is_correct:
                                correct_count += 1
                        
                        score_percent = (submission.get("score", 0) / max_possible) * 100 if max_possible > 0 else 0
                        
                        # Thêm dữ liệu vào danh sách
                        entry = {
                            "Lần làm": idx + 1,
                            "Thời gian": submission_time,
                            "Điểm số": submission.get("score", 0),
                            "Điểm tối đa": max_possible,
                            "Tỷ lệ đúng": f"{score_percent:.1f}%",
                            "Số câu đúng": f"{correct_count}/{len(questions)}"
                        }
                        
                        # Thêm chi tiết từng câu hỏi
                        for q in questions:
                            q_id = str(q.get("id", ""))
                            user_ans = responses.get(q_id, [])
                            is_correct = check_answer_correctness(user_ans, q)
                            
                            entry[f"Câu {q_id}"] = ", ".join([str(a) for a in user_ans]) if user_ans else "Không trả lời"
                            entry[f"Câu {q_id} - Kết quả"] = "Đúng" if is_correct else "Sai"
                        
                        student_report_data.append(entry)
                    
                    # Tạo DataFrame cho báo cáo
                    df_student_report = pd.DataFrame(student_report_data)
                    
                    # Hiển thị dữ liệu dạng bảng
                    st.write("### Chi tiết các lần làm bài")
                    st.dataframe(df_student_report, hide_index=True, use_container_width=True)
                    
                    # Tạo báo cáo Word cho học viên này
                    try:
                        # Tạo tiêu đề
                        title = f"Báo cáo chi tiết học viên: {student_name} ({selected_email})"
                        
                        st.write("### Tải xuống báo cáo học viên")
                        
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            # Word
                            docx_buffer = dataframe_to_docx(df_student_report, title, f"bao_cao_{student_name}.docx")
                            st.markdown(
                                get_download_link_docx(docx_buffer, 
                                                    f"bao_cao_{student_name.replace(' ', '_')}.docx", 
                                                    "Tải xuống báo cáo DOCX"), 
                                unsafe_allow_html=True
                            )
                        
                        with col2:
                            # PDF
                            pdf_buffer = dataframe_to_pdf_fpdf(df_student_report, title, f"bao_cao_{student_name}.pdf")
                            st.markdown(
                                get_download_link_pdf(pdf_buffer, 
                                                    f"bao_cao_{student_name.replace(' ', '_')}.pdf", 
                                                    "Tải xuống báo cáo PDF"), 
                                unsafe_allow_html=True
                            )
                        
                        # Tạo báo cáo chi tiết cho từng lần làm
                        st.write("### Tải báo cáo chi tiết từng lần làm")
                        for idx, submission in enumerate(student_submissions):
                            col1, col2 = st.columns(2)
                            
                            with col1:
                                try:
                                    # Word
                                    docx_buffer = create_student_report_docx(
                                        student_name,
                                        selected_email,
                                        student_class,
                                        submission,
                                        questions,
                                        max_possible
                                    )
                                    
                                    st.markdown(
                                        get_download_link_docx(
                                            docx_buffer, 
                                            f"bao_cao_chi_tiet_{student_name.replace(' ', '_')}_lan_{idx+1}.docx", 
                                            f"Tải xuống báo cáo lần {idx+1} (DOCX)"
                                        ), 
                                        unsafe_allow_html=True
                                    )
                                except Exception as e:
                                    st.error(f"Lỗi khi tạo báo cáo DOCX lần {idx+1}: {str(e)}")
                            
                            with col2:
                                try:
                                    # PDF
                                    pdf_buffer = create_student_report_pdf_fpdf(
                                        student_name,
                                        selected_email,
                                        student_class,
                                        submission,
                                        questions,
                                        max_possible
                                    )
                                    
                                    st.markdown(
                                        get_download_link_pdf(
                                            pdf_buffer, 
                                            f"bao_cao_chi_tiet_{student_name.replace(' ', '_')}_lan_{idx+1}.pdf", 
                                            f"Tải xuống báo cáo lần {idx+1} (PDF)"
                                        ), 
                                        unsafe_allow_html=True
                                    )
                                except Exception as e:
                                    st.error(f"Lỗi khi tạo báo cáo PDF lần {idx+1}: {str(e)}")
                        
                    except Exception as e:
                        st.error(f"Lỗi khi tạo báo cáo: {str(e)}")
                
                else:
                    st.warning(f"Không tìm thấy bài nộp nào của học viên {student_name} ({selected_email})")
            else:
                st.info("Vui lòng chọn email học viên để xem và xuất báo cáo")
                
        except Exception as e:
            st.error(f"Lỗi khi xử lý báo cáo theo học viên: {str(e)}")

def view_statistics():
    """Hiển thị trang thống kê và báo cáo"""
    st.title("📊 Báo cáo & thống kê")
    
    # Khởi tạo biến trước
    questions = []
    students = []
    submissions = []
    max_possible = 0
    df_questions = pd.DataFrame()
    df_students_list = pd.DataFrame()
    df_class_stats = pd.DataFrame()
    df_all_submissions = pd.DataFrame()
    
    try:
        # Lấy dữ liệu từ database
        questions = get_all_questions()
        students = get_all_users(role="Học viên")
        
        # Tạo form tìm kiếm email nếu muốn xem báo cáo theo học viên cụ thể
        with st.sidebar:
            st.subheader("Tìm kiếm học viên")
            search_email = st.text_input("Nhập email học viên:", key="search_email_stats")
            search_button = st.button("Tìm kiếm", key="search_button_stats")
        
        if search_button and search_email:
            submissions = get_user_submissions(search_email)
            if not submissions:
                st.warning(f"Không tìm thấy bài nộp của học viên: {search_email}")
                return
        else:
            # Lấy tất cả bài nộp từ tất cả học viên
            for student in students:
                try:
                    student_submissions = get_user_submissions(student.get("email", ""))
                    submissions.extend(student_submissions)
                except Exception as e:
                    st.error(f"Lỗi khi lấy dữ liệu của học viên {student.get('email', '')}: {str(e)}")
        
        if not questions:
            st.warning("Chưa có dữ liệu câu hỏi nào trong hệ thống.")
            return
        
        if not submissions:
            st.warning("Chưa có ai nộp khảo sát.")
            return
        
        # Tạo tab thống kê
        tab1, tab2, tab3, tab4, tab5 = st.tabs(["Tổng quan", "Theo học viên", "Theo câu hỏi", "Danh sách học viên", "Xuất báo cáo"])
        
        # Tính tổng điểm tối đa
        max_possible = sum([q.get("score", 0) for q in questions])
        
        # Chuẩn bị dữ liệu cho tất cả các submissions
        all_submission_data = []
        
        for s in submissions:
            try:
                # Tìm thông tin học viên
                student_info = next((student for student in students if student.get("email") == s.get("user_email")), None)
                full_name = student_info.get("full_name", "Không xác định") if student_info else "Không xác định"
                class_name = student_info.get("class", "Không xác định") if student_info else "Không xác định"
                
                # Chuyển đổi timestamp sang định dạng đọc được
                submission_time = "Không xác định"
                if isinstance(s.get("timestamp"), (int, float)):
                    try:
                        submission_time = datetime.fromtimestamp(s.get("timestamp")).strftime("%d/%m/%Y %H:%M:%S")
                    except:
                        pass
                else:
                    try:
                        dt = datetime.fromisoformat(s.get("timestamp", "").replace("Z", "+00:00"))
                        submission_time = dt.strftime("%d/%m/%Y %H:%M:%S")
                    except:
                        pass
                
                # Thêm thông tin cơ bản
                submission_data = {
                    "ID": s.get("id", ""),
                    "Email": s.get("user_email", ""),
                    "Họ và tên": full_name,
                    "Lớp": class_name,
                    "Thời gian nộp": submission_time,
                    "Điểm số": s.get("score", 0),
                    "Điểm tối đa": max_possible,
                    "Tỷ lệ đúng": f"{(s.get('score', 0)/max_possible*100):.1f}%" if max_possible > 0 else "N/A"
                }
                
                # Chuyển đổi responses từ JSON string thành dict nếu cần
                responses = s.get("responses", {})
                if isinstance(responses, str):
                    try:
                        responses = json.loads(responses)
                    except:
                        responses = {}
                
                # Thêm câu trả lời của từng câu hỏi
                for q in questions:
                    q_id = str(q.get("id", ""))
                    user_ans = responses.get(q_id, [])
                    
                    # Đảm bảo q["correct"] và q["answers"] có định dạng đúng
                    q_correct = q.get("correct", [])
                    q_answers = q.get("answers", [])
                    
                    if isinstance(q_correct, str):
                        try:
                            q_correct = json.loads(q_correct)
                        except:
                            try:
                                q_correct = [int(x.strip()) for x in q_correct.split(",")]
                            except:
                                q_correct = []
                    
                    if isinstance(q_answers, str):
                        try:
                            q_answers = json.loads(q_answers)
                        except:
                            q_answers = [q_answers]
                    
                    try:
                        expected = [q_answers[i - 1] for i in q_correct]
                    except (IndexError, TypeError):
                        expected = ["Lỗi đáp án"]
                        
                    is_correct = check_answer_correctness(user_ans, q)
                    
                    # Thêm thông tin câu hỏi
                    submission_data[f"Câu {q_id}: {q.get('question', '')}"] = ", ".join([str(a) for a in user_ans]) if user_ans else "Không trả lời"
                    submission_data[f"Câu {q_id} - Đúng/Sai"] = "Đúng" if is_correct else "Sai"
                
                all_submission_data.append(submission_data)
            except Exception as e:
                st.error(f"Lỗi khi xử lý submission ID {s.get('id', '')}: {str(e)}")
        
        # DataFrame chứa tất cả bài nộp
        df_all_submissions = pd.DataFrame(all_submission_data) if all_submission_data else pd.DataFrame()
        
        with tab1:
            display_overview_tab(submissions, students, questions, max_possible)
        
        with tab2:
            display_student_tab(submissions, students, questions, max_possible)
        
        with tab3:
            df_questions = display_question_tab(submissions, questions)
        
        with tab4:
            df_students_list, df_class_stats = display_student_list_tab(submissions, students, max_possible)
        
        with tab5:
            display_export_tab(df_all_submissions, df_questions, df_students_list, df_class_stats)
    
    except Exception as e:
        st.error(f"Đã xảy ra lỗi không mong muốn: {str(e)}")
        traceback.print_exc()



# Chỉ chạy hàm main khi chạy file này trực tiếp
if __name__ == "__main__":
    st.set_page_config(
        page_title="Báo cáo & Thống kê",
        page_icon="📊",
        layout="wide",
    )
    view_statistics()
